"""
voxnotes.py
PowerPoint speaker notes export/import for VoxPrep.

Extracts speaker notes from PowerPoint decks and exports to Word/text/markdown.
Imports edited notes back into PowerPoint with change detection.

Export formats:
    - .docx: VO-friendly Word document (sans-serif, 14pt, generous spacing)
    - .txt: Plain text with slide markers
    - .md: Markdown with slide headers

Example:
    notes = extract_notes("Training.pptx")
    export_to_docx(notes, "Training_notes.docx")
    
    # After editing...
    edited = parse_notes_file("Training_notes.docx")
    changes = compare_notes(notes, edited)
    apply_notes("Training.pptx", changes)
"""

import os
import re
import time
from pathlib import Path
from typing import List, Dict, Optional, Callable

try:
    from win32com.client import Dispatch, gencache
    HAS_COM = True
except Exception:
    HAS_COM = False

try:
    import win32api
    HAS_WIN32API = True
except Exception:
    HAS_WIN32API = False

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_LINE_SPACING
    HAS_DOCX = True
except Exception:
    HAS_DOCX = False

LOG_PREFIX = "[voxnotes]"

# COM retry settings
COM_RETRY_ATTEMPTS = 3
COM_RETRY_DELAY = 1.5  # seconds


def log(msg: str):
    """Simple logging helper."""
    print(f"{LOG_PREFIX} {msg}", flush=True)


def open_presentation_with_retry(app, pptx_path: str, read_only: bool = True, max_attempts: int = COM_RETRY_ATTEMPTS):
    """
    Open a PowerPoint presentation with retry logic.
    
    PowerPoint COM can fail if a previous instance hasn't fully released.
    This retries with a delay to handle that race condition.
    """
    last_error = None
    
    for attempt in range(1, max_attempts + 1):
        try:
            pres = app.Presentations.Open(pptx_path, read_only, False, False)
            return pres
        except Exception as e:
            last_error = e
            if attempt < max_attempts:
                log(f"  Open failed (attempt {attempt}/{max_attempts}), retrying in {COM_RETRY_DELAY}s...")
                time.sleep(COM_RETRY_DELAY)
            else:
                log(f"  Open failed after {max_attempts} attempts")
    
    raise RuntimeError(f"Failed to open presentation after {max_attempts} attempts: {last_error}")


def sanitize_text(text: str) -> str:
    """
    Remove control characters that break XML/docx.
    
    Keeps: tab (0x09), newline (0x0A), carriage return (0x0D)
    Removes: NULL bytes, other control chars (0x00-0x08, 0x0B-0x0C, 0x0E-0x1F)
    
    Args:
        text: String that may contain control characters
    
    Returns:
        Clean string safe for XML
    """
    if not text:
        return text
    
    # Build translation table: control chars -> None (delete)
    # Keep: 0x09 (tab), 0x0A (LF), 0x0D (CR)
    delete_chars = ''.join(
        chr(i) for i in range(32) 
        if i not in (9, 10, 13)  # Keep tab, LF, CR
    )
    # Also remove DEL (0x7F) and other problematic chars
    delete_chars += chr(127)
    
    return text.translate(str.maketrans('', '', delete_chars))


def get_short_path(long_path: str) -> str:
    """
    Convert long path with spaces to Windows 8.3 short path format.
    PowerPoint COM doesn't like spaces in paths.
    """
    if not HAS_WIN32API:
        return long_path
    
    try:
        long_path = long_path.replace('/', '\\')
        short_path = win32api.GetShortPathName(long_path)
        return short_path
    except Exception:
        return long_path


# =============================================================================
# EXTRACTION
# =============================================================================

def extract_notes(pptx_path: str, log_callback: Callable = None) -> List[Dict]:
    """
    Extract speaker notes from all slides in a PowerPoint deck.
    
    Args:
        pptx_path: Path to .pptx file
        log_callback: Optional function for progress logging
    
    Returns:
        List of dicts, one per slide:
        [
            {
                "slide_number": 1,
                "slide_title": "Introduction",
                "notes": "Welcome to this training module..."
            },
            ...
        ]
    
    Raises:
        RuntimeError: If COM not available or file can't be opened
    """
    def _log(msg):
        if log_callback:
            log_callback(msg)
        else:
            log(msg)
    
    if not HAS_COM:
        raise RuntimeError("Windows COM API not available (pywin32 not installed)")
    
    pptx_path = get_short_path(str(Path(pptx_path).resolve()))
    
    if not os.path.isfile(pptx_path):
        raise FileNotFoundError(f"PowerPoint file not found: {pptx_path}")
    
    app = None
    pres = None
    
    try:
        gencache.EnsureDispatch("PowerPoint.Application")
        app = Dispatch("PowerPoint.Application")
        app.Visible = True
        
        pres = open_presentation_with_retry(app, pptx_path, read_only=True)
        
        notes_data = []
        slide_count = pres.Slides.Count
        
        _log(f"Extracting notes from {slide_count} slides...")
        
        for i in range(1, slide_count + 1):
            slide = pres.Slides.Item(i)
            
            # Get slide title (first title shape or placeholder)
            title = ""
            try:
                for shape in slide.Shapes:
                    if shape.HasTextFrame:
                        if shape.PlaceholderFormat.Type == 1:  # ppPlaceholderTitle
                            title = shape.TextFrame.TextRange.Text.strip()
                            break
            except:
                pass
            
            # If no title placeholder, try first text shape
            if not title:
                try:
                    for shape in slide.Shapes:
                        if shape.HasTextFrame and shape.TextFrame.HasText:
                            text = shape.TextFrame.TextRange.Text.strip()
                            if text and len(text) < 100:  # Reasonable title length
                                title = text
                                break
                except:
                    pass
            
            # Get speaker notes
            notes_text = ""
            try:
                notes_page = slide.NotesPage
                for shape in notes_page.Shapes:
                    if shape.PlaceholderFormat.Type == 2:  # ppPlaceholderBody
                        if shape.HasTextFrame:
                            notes_text = shape.TextFrame.TextRange.Text.strip()
                            break
            except:
                pass
            
            notes_data.append({
                "slide_number": i,
                "slide_title": sanitize_text(title),
                "notes": sanitize_text(notes_text)
            })
            
            if notes_text:
                _log(f"  Slide {i}: {len(notes_text)} chars")
            else:
                _log(f"  Slide {i}: (no notes)")
        
        notes_count = sum(1 for n in notes_data if n["notes"])
        _log(f"Extracted notes from {notes_count} of {slide_count} slides")
        
        return notes_data
        
    except Exception as e:
        raise RuntimeError(f"Failed to extract notes: {e}")
    
    finally:
        # Always clean up COM objects
        if pres is not None:
            try:
                pres.Close()
            except:
                pass
        if app is not None:
            try:
                app.Quit()
            except:
                pass


# =============================================================================
# EXPORT FUNCTIONS
# =============================================================================

def export_to_docx(notes: List[Dict], output_path: str, 
                   font_name: str = "Calibri", font_size: int = 14,
                   log_callback: Callable = None) -> str:
    """
    Export notes to Word document formatted for VO recording.
    
    Args:
        notes: List of note dicts from extract_notes()
        output_path: Where to save .docx file
        font_name: Font family (default: Calibri)
        font_size: Font size in points (default: 14)
        log_callback: Optional function for progress logging
    
    Returns:
        Path to created file
    
    Format:
        - Sans-serif font (Calibri/Arial)
        - 14pt or larger
        - 1.5 line spacing
        - Clear slide separators
        - NO thumbnails (what breaks PowerPoint's export)
    """
    def _log(msg):
        if log_callback:
            log_callback(msg)
        else:
            log(msg)
    
    if not HAS_DOCX:
        raise RuntimeError("python-docx not installed (pip install python-docx)")
    
    _log(f"Creating Word document: {os.path.basename(output_path)}")
    
    doc = Document()
    
    # Set default font for document
    style = doc.styles['Normal']
    style.font.name = font_name
    style.font.size = Pt(font_size)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    style.paragraph_format.space_after = Pt(12)
    
    slides_with_notes = 0
    
    for note in notes:
        slide_num = note["slide_number"]
        title = note["slide_title"]
        notes_text = note["notes"]
        
        # Slide header
        if title:
            header_text = f"Slide {slide_num}: {sanitize_text(title)}"
        else:
            header_text = f"Slide {slide_num}"
        
        # Add slide header (bold, slightly larger)
        header = doc.add_paragraph()
        header_run = header.add_run(header_text)
        header_run.bold = True
        header_run.font.size = Pt(font_size + 2)
        header.paragraph_format.space_before = Pt(18)
        header.paragraph_format.space_after = Pt(6)
        
        # Add separator line
        separator = doc.add_paragraph()
        sep_run = separator.add_run("─" * 50)
        sep_run.font.size = Pt(10)
        separator.paragraph_format.space_after = Pt(12)
        
        # Add notes content (or placeholder)
        if notes_text:
            # Sanitize and split by paragraphs
            clean_notes = sanitize_text(notes_text)
            paragraphs = clean_notes.split('\n')
            for para_text in paragraphs:
                if para_text.strip():
                    para = doc.add_paragraph(para_text.strip())
                    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            slides_with_notes += 1
        else:
            # Empty placeholder for slides without notes
            para = doc.add_paragraph("[No notes]")
            para.runs[0].italic = True
        
        # Add spacing after each slide's notes
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(6)
    
    doc.save(output_path)
    
    _log(f"Exported {slides_with_notes} slides with notes to {os.path.basename(output_path)}")
    
    return output_path


def export_to_txt(notes: List[Dict], output_path: str,
                  log_callback: Callable = None) -> str:
    """
    Export notes to plain text file.
    
    Args:
        notes: List of note dicts from extract_notes()
        output_path: Where to save .txt file
        log_callback: Optional function for progress logging
    
    Returns:
        Path to created file
    
    Format:
        Slide 1: Introduction
        ──────────────────────────────────────────────────
        
        Welcome to this training module...
        
        ══════════════════════════════════════════════════
        
        Slide 2: Prerequisites
        ...
    """
    def _log(msg):
        if log_callback:
            log_callback(msg)
        else:
            log(msg)
    
    _log(f"Creating text file: {os.path.basename(output_path)}")
    
    lines = []
    slides_with_notes = 0
    
    for note in notes:
        slide_num = note["slide_number"]
        title = note["slide_title"]
        notes_text = note["notes"]
        
        # Slide header
        if title:
            lines.append(f"Slide {slide_num}: {title}")
        else:
            lines.append(f"Slide {slide_num}")
        
        lines.append("─" * 50)
        lines.append("")
        
        # Notes content
        if notes_text:
            lines.append(notes_text)
            slides_with_notes += 1
        else:
            lines.append("[No notes]")
        
        lines.append("")
        lines.append("═" * 50)
        lines.append("")
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))
    
    _log(f"Exported {slides_with_notes} slides with notes to {os.path.basename(output_path)}")
    
    return output_path


def export_to_md(notes: List[Dict], output_path: str,
                 log_callback: Callable = None) -> str:
    """
    Export notes to Markdown file.
    
    Args:
        notes: List of note dicts from extract_notes()
        output_path: Where to save .md file
        log_callback: Optional function for progress logging
    
    Returns:
        Path to created file
    
    Format:
        ## Slide 1: Introduction
        
        Welcome to this training module...
        
        ---
        
        ## Slide 2: Prerequisites
        ...
    """
    def _log(msg):
        if log_callback:
            log_callback(msg)
        else:
            log(msg)
    
    _log(f"Creating Markdown file: {os.path.basename(output_path)}")
    
    lines = []
    slides_with_notes = 0
    
    # Document title
    lines.append("# Speaker Notes")
    lines.append("")
    
    for note in notes:
        slide_num = note["slide_number"]
        title = note["slide_title"]
        notes_text = note["notes"]
        
        # Slide header (H2)
        if title:
            lines.append(f"## Slide {slide_num}: {title}")
        else:
            lines.append(f"## Slide {slide_num}")
        
        lines.append("")
        
        # Notes content
        if notes_text:
            lines.append(notes_text)
            slides_with_notes += 1
        else:
            lines.append("*[No notes]*")
        
        lines.append("")
        lines.append("---")
        lines.append("")
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))
    
    _log(f"Exported {slides_with_notes} slides with notes to {os.path.basename(output_path)}")
    
    return output_path


# =============================================================================
# CONVENIENCE FUNCTION
# =============================================================================

def export_notes(pptx_path: str, output_path: str, format: str = "docx",
                 font_name: str = "Calibri", font_size: int = 14,
                 log_callback: Callable = None) -> str:
    """
    Convenience function: extract and export notes in one call.
    
    Args:
        pptx_path: Path to PowerPoint file
        output_path: Where to save exported notes
        format: "docx", "txt", or "md"
        font_name: Font for docx export (default: Calibri)
        font_size: Font size for docx export (default: 14)
        log_callback: Optional function for progress logging
    
    Returns:
        Path to created file
    
    Example:
        export_notes("Training.pptx", "Training_notes.docx", "docx")
    """
    def _log(msg):
        if log_callback:
            log_callback(msg)
        else:
            log(msg)
    
    # Extract notes
    notes = extract_notes(pptx_path, log_callback)
    
    # Export based on format
    format = format.lower().strip('.')
    
    if format == "docx":
        return export_to_docx(notes, output_path, font_name, font_size, log_callback)
    elif format == "txt":
        return export_to_txt(notes, output_path, log_callback)
    elif format == "md":
        return export_to_md(notes, output_path, log_callback)
    else:
        raise ValueError(f"Unknown format: {format}. Use 'docx', 'txt', or 'md'.")


# =============================================================================
# IMPORT FUNCTIONS (Parse edited files, compare, apply changes)
# =============================================================================

def parse_notes_file(file_path: str, log_callback: Callable = None) -> List[Dict]:
    """
    Parse an edited notes file (docx, txt, or md) back into note dicts.
    
    Args:
        file_path: Path to edited notes file
        log_callback: Optional function for progress logging
    
    Returns:
        List of dicts matching extract_notes() format:
        [{"slide_number": 1, "slide_title": "...", "notes": "..."}, ...]
    
    Raises:
        ValueError: If file format not recognized or parsing fails
    """
    def _log(msg):
        if log_callback:
            log_callback(msg)
        else:
            log(msg)
    
    ext = Path(file_path).suffix.lower()
    
    if ext == '.docx':
        return _parse_docx(file_path, log_callback)
    elif ext == '.txt':
        return _parse_txt(file_path, log_callback)
    elif ext == '.md':
        return _parse_md(file_path, log_callback)
    else:
        raise ValueError(f"Unknown file format: {ext}. Use .docx, .txt, or .md")


def _parse_docx(file_path: str, log_callback: Callable = None) -> List[Dict]:
    """Parse notes from edited Word document."""
    def _log(msg):
        if log_callback:
            log_callback(msg)
        else:
            log(msg)
    
    if not HAS_DOCX:
        raise RuntimeError("python-docx not installed")
    
    _log(f"Parsing Word document: {os.path.basename(file_path)}")
    
    doc = Document(file_path)
    
    notes_data = []
    current_slide = None
    current_notes = []
    
    # Pattern to match slide headers: "Slide N" or "Slide N: Title"
    slide_pattern = re.compile(r'^Slide\s+(\d+)(?::\s*(.*))?$', re.IGNORECASE)
    separator_pattern = re.compile(r'^[─═]{10,}$')
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # Skip empty lines and separators
        if not text or separator_pattern.match(text):
            continue
        
        # Check for slide header
        match = slide_pattern.match(text)
        if match:
            # Save previous slide if exists
            if current_slide is not None:
                notes_text = '\n'.join(current_notes).strip()
                # Remove placeholder text
                if notes_text == "[No notes]":
                    notes_text = ""
                notes_data.append({
                    "slide_number": current_slide["number"],
                    "slide_title": current_slide["title"],
                    "notes": notes_text
                })
            
            # Start new slide
            current_slide = {
                "number": int(match.group(1)),
                "title": match.group(2) or ""
            }
            current_notes = []
        elif current_slide is not None:
            # Add to current slide's notes (skip [No notes] placeholder)
            if text != "[No notes]":
                current_notes.append(text)
    
    # Don't forget last slide
    if current_slide is not None:
        notes_text = '\n'.join(current_notes).strip()
        if notes_text == "[No notes]":
            notes_text = ""
        notes_data.append({
            "slide_number": current_slide["number"],
            "slide_title": current_slide["title"],
            "notes": notes_text
        })
    
    _log(f"Parsed {len(notes_data)} slides from document")
    
    return notes_data


def _parse_txt(file_path: str, log_callback: Callable = None) -> List[Dict]:
    """Parse notes from edited text file."""
    def _log(msg):
        if log_callback:
            log_callback(msg)
        else:
            log(msg)
    
    _log(f"Parsing text file: {os.path.basename(file_path)}")
    
    with open(file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    notes_data = []
    current_slide = None
    current_notes = []
    
    slide_pattern = re.compile(r'^Slide\s+(\d+)(?::\s*(.*))?$', re.IGNORECASE)
    separator_pattern = re.compile(r'^[─═]{10,}$')
    
    for line in content.split('\n'):
        text = line.strip()
        
        # Skip separators
        if separator_pattern.match(text):
            continue
        
        # Check for slide header
        match = slide_pattern.match(text)
        if match:
            # Save previous slide
            if current_slide is not None:
                notes_text = '\n'.join(current_notes).strip()
                if notes_text == "[No notes]":
                    notes_text = ""
                notes_data.append({
                    "slide_number": current_slide["number"],
                    "slide_title": current_slide["title"],
                    "notes": notes_text
                })
            
            current_slide = {
                "number": int(match.group(1)),
                "title": match.group(2) or ""
            }
            current_notes = []
        elif current_slide is not None:
            # Add to notes (preserve empty lines within notes)
            if text != "[No notes]":
                current_notes.append(line.rstrip())  # Keep original spacing
    
    # Last slide
    if current_slide is not None:
        notes_text = '\n'.join(current_notes).strip()
        if notes_text == "[No notes]":
            notes_text = ""
        notes_data.append({
            "slide_number": current_slide["number"],
            "slide_title": current_slide["title"],
            "notes": notes_text
        })
    
    _log(f"Parsed {len(notes_data)} slides from text file")
    
    return notes_data


def _parse_md(file_path: str, log_callback: Callable = None) -> List[Dict]:
    """Parse notes from edited Markdown file."""
    def _log(msg):
        if log_callback:
            log_callback(msg)
        else:
            log(msg)
    
    _log(f"Parsing Markdown file: {os.path.basename(file_path)}")
    
    with open(file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    notes_data = []
    current_slide = None
    current_notes = []
    
    # Markdown headers: ## Slide N or ## Slide N: Title
    slide_pattern = re.compile(r'^##\s+Slide\s+(\d+)(?::\s*(.*))?$', re.IGNORECASE)
    
    for line in content.split('\n'):
        text = line.strip()
        
        # Skip document title and horizontal rules
        if text.startswith('# ') or text == '---':
            continue
        
        # Check for slide header
        match = slide_pattern.match(text)
        if match:
            # Save previous slide
            if current_slide is not None:
                notes_text = '\n'.join(current_notes).strip()
                # Remove markdown italic placeholder
                if notes_text == "*[No notes]*":
                    notes_text = ""
                notes_data.append({
                    "slide_number": current_slide["number"],
                    "slide_title": current_slide["title"],
                    "notes": notes_text
                })
            
            current_slide = {
                "number": int(match.group(1)),
                "title": match.group(2) or ""
            }
            current_notes = []
        elif current_slide is not None:
            if text != "*[No notes]*":
                current_notes.append(line.rstrip())
    
    # Last slide
    if current_slide is not None:
        notes_text = '\n'.join(current_notes).strip()
        if notes_text == "*[No notes]*":
            notes_text = ""
        notes_data.append({
            "slide_number": current_slide["number"],
            "slide_title": current_slide["title"],
            "notes": notes_text
        })
    
    _log(f"Parsed {len(notes_data)} slides from Markdown file")
    
    return notes_data


def compare_notes(original: List[Dict], edited: List[Dict],
                  log_callback: Callable = None) -> List[Dict]:
    """
    Compare original notes with edited version to find changes.
    
    Args:
        original: Notes from extract_notes()
        edited: Notes from parse_notes_file()
        log_callback: Optional function for progress logging
    
    Returns:
        List of changes:
        [
            {
                "slide_number": 3,
                "slide_title": "...",
                "original_notes": "old text...",
                "edited_notes": "new text...",
                "change_type": "modified"  # "modified", "added", "removed"
            },
            ...
        ]
    """
    def _log(msg):
        if log_callback:
            log_callback(msg)
        else:
            log(msg)
    
    # Build lookup by slide number
    original_by_slide = {n["slide_number"]: n for n in original}
    edited_by_slide = {n["slide_number"]: n for n in edited}
    
    changes = []
    
    # Check all slides in edited version
    for slide_num, edited_note in edited_by_slide.items():
        original_note = original_by_slide.get(slide_num)
        
        if original_note is None:
            # New slide (shouldn't happen normally, but handle it)
            if edited_note["notes"]:
                changes.append({
                    "slide_number": slide_num,
                    "slide_title": edited_note["slide_title"],
                    "original_notes": "",
                    "edited_notes": edited_note["notes"],
                    "change_type": "added"
                })
        else:
            # Compare notes
            orig_text = original_note["notes"].strip()
            edit_text = edited_note["notes"].strip()
            
            if orig_text != edit_text:
                if not orig_text and edit_text:
                    change_type = "added"
                elif orig_text and not edit_text:
                    change_type = "removed"
                else:
                    change_type = "modified"
                
                changes.append({
                    "slide_number": slide_num,
                    "slide_title": original_note["slide_title"],
                    "original_notes": orig_text,
                    "edited_notes": edit_text,
                    "change_type": change_type
                })
    
    _log(f"Found {len(changes)} changed slide(s)")
    
    return changes


def apply_notes(pptx_path: str, changes: List[Dict], 
                slides_to_apply: List[int] = None,
                log_callback: Callable = None) -> Dict:
    """
    Apply note changes back to PowerPoint file.
    
    Args:
        pptx_path: Path to PowerPoint file
        changes: List of changes from compare_notes()
        slides_to_apply: Optional list of slide numbers to apply (default: all)
        log_callback: Optional function for progress logging
    
    Returns:
        Result dict:
        {
            "applied": [3, 7, 12],  # Slides that were updated
            "skipped": [5],         # Slides that were skipped
            "errors": []            # Any errors
        }
    """
    def _log(msg):
        if log_callback:
            log_callback(msg)
        else:
            log(msg)
    
    if not HAS_COM:
        raise RuntimeError("Windows COM API not available")
    
    pptx_path = get_short_path(str(Path(pptx_path).resolve()))
    
    # Filter changes if specific slides requested
    if slides_to_apply is not None:
        changes = [c for c in changes if c["slide_number"] in slides_to_apply]
    
    if not changes:
        _log("No changes to apply")
        return {"applied": [], "skipped": [], "errors": []}
    
    _log(f"Applying changes to {len(changes)} slide(s)...")
    
    result = {"applied": [], "skipped": [], "errors": []}
    
    app = None
    pres = None
    
    try:
        gencache.EnsureDispatch("PowerPoint.Application")
        app = Dispatch("PowerPoint.Application")
        app.Visible = True
        
        # Open for editing (not read-only) with retry
        pres = open_presentation_with_retry(app, pptx_path, read_only=False)
        
        for change in changes:
            slide_num = change["slide_number"]
            new_notes = change["edited_notes"]
            
            try:
                slide = pres.Slides.Item(slide_num)
                notes_page = slide.NotesPage
                
                # Find notes placeholder
                for shape in notes_page.Shapes:
                    try:
                        if shape.PlaceholderFormat.Type == 2:  # ppPlaceholderBody
                            if shape.HasTextFrame:
                                shape.TextFrame.TextRange.Text = new_notes
                                result["applied"].append(slide_num)
                                _log(f"  Slide {slide_num}: Updated")
                                break
                    except:
                        continue
                        
            except Exception as e:
                result["errors"].append(f"Slide {slide_num}: {e}")
                _log(f"  Slide {slide_num}: Error - {e}")
        
        # Save changes
        pres.Save()
        
        _log(f"Applied changes to {len(result['applied'])} slide(s)")
        
        return result
        
    except Exception as e:
        raise RuntimeError(f"Failed to apply notes: {e}")
    
    finally:
        # Always clean up COM objects
        if pres is not None:
            try:
                pres.Close()
            except:
                pass
        if app is not None:
            try:
                app.Quit()
            except:
                pass


def import_notes(pptx_path: str, notes_file: str, preview_only: bool = False,
                 slides_to_apply: List[int] = None,
                 log_callback: Callable = None) -> Dict:
    """
    Convenience function: parse, compare, and optionally apply changes.
    
    Args:
        pptx_path: Path to PowerPoint file
        notes_file: Path to edited notes file (docx, txt, or md)
        preview_only: If True, return changes without applying
        slides_to_apply: Optional list of slide numbers to apply
        log_callback: Optional function for progress logging
    
    Returns:
        If preview_only:
            {"changes": [...], "preview": True}
        If applying:
            {"applied": [...], "skipped": [...], "errors": [...]}
    
    Example:
        # Preview changes
        result = import_notes("Training.pptx", "Training_notes.docx", preview_only=True)
        print(f"Found {len(result['changes'])} changes")
        
        # Apply all changes
        result = import_notes("Training.pptx", "Training_notes.docx")
        print(f"Applied to {len(result['applied'])} slides")
    """
    def _log(msg):
        if log_callback:
            log_callback(msg)
        else:
            log(msg)
    
    # Extract original notes
    original = extract_notes(pptx_path, log_callback)
    
    # Parse edited file
    edited = parse_notes_file(notes_file, log_callback)
    
    # Compare
    changes = compare_notes(original, edited, log_callback)
    
    if preview_only:
        return {"changes": changes, "preview": True}
    
    # Apply changes
    if not changes:
        return {"applied": [], "skipped": [], "errors": []}
    
    return apply_notes(pptx_path, changes, slides_to_apply, log_callback)


# =============================================================================
# CLI for standalone testing
# =============================================================================

def _usage():
    print("Usage: python voxnotes.py <command> <args>")
    print()
    print("Commands:")
    print("  export <deck.pptx> <output_file>     Export notes to file")
    print("  import <deck.pptx> <notes_file>      Import notes from file")
    print("  preview <deck.pptx> <notes_file>     Preview changes without applying")
    print()
    print("Output formats: .docx, .txt, .md (determined by extension)")


if __name__ == "__main__":
    import sys
    
    if len(sys.argv) < 2:
        _usage()
        sys.exit(64)
    
    command = sys.argv[1].lower()
    
    try:
        if command == "export":
            if len(sys.argv) < 4:
                print("Usage: python voxnotes.py export <deck.pptx> <output_file>")
                sys.exit(64)
            
            pptx_path = sys.argv[2]
            output_path = sys.argv[3]
            ext = Path(output_path).suffix.lower().strip('.')
            
            result = export_notes(pptx_path, output_path, ext)
            print(f"\nCreated: {result}")
            
        elif command == "import":
            if len(sys.argv) < 4:
                print("Usage: python voxnotes.py import <deck.pptx> <notes_file>")
                sys.exit(64)
            
            pptx_path = sys.argv[2]
            notes_file = sys.argv[3]
            
            result = import_notes(pptx_path, notes_file)
            print(f"\nApplied to slides: {result['applied']}")
            if result['errors']:
                print(f"Errors: {result['errors']}")
                
        elif command == "preview":
            if len(sys.argv) < 4:
                print("Usage: python voxnotes.py preview <deck.pptx> <notes_file>")
                sys.exit(64)
            
            pptx_path = sys.argv[2]
            notes_file = sys.argv[3]
            
            result = import_notes(pptx_path, notes_file, preview_only=True)
            
            if not result['changes']:
                print("\nNo changes detected.")
            else:
                print(f"\nFound {len(result['changes'])} change(s):")
                for change in result['changes']:
                    print(f"  Slide {change['slide_number']}: {change['change_type']}")
                    
        else:
            print(f"Unknown command: {command}")
            _usage()
            sys.exit(64)
            
    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(2)
